# -*- coding: utf-8 -*-
import sys
import time
import os
# 用于pdf文件的处理
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfpage import PDFTextExtractionNotAllowed
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.pdfdevice import PDFDevice
from pdfminer.layout import *
from pdfminer.converter import PDFPageAggregator

###########################################把pdf转为txt################################################
def p2t(sourcefile, outfile):
    with open(sourcefile, 'rb') as fp:
        # 来创建一个pdf文档分析器
        parser = PDFParser(fp)
        #创建一个PDF文档对象存储文档结构
        try:
            document = PDFDocument(parser)
        except:
            print(sourcefile + ' :pdf未正确下载')
        # 检查文件是否允许文本提取
        else:
            if not document.is_extractable:
                print(sourcefile + ' :不允许提取文本')
             # 创建一个PDF资源管理器对象来存储共赏资源
            rsrcmgr=PDFResourceManager()
             # 设定参数进行分析
            laparams=LAParams()
             # 创建一个PDF设备对象
             # device=PDFDevice(rsrcmgr)
            device=PDFPageAggregator(rsrcmgr,laparams=laparams)
             # 创建一个PDF解释器对象
            interpreter=PDFPageInterpreter(rsrcmgr,device)
             # 处理每一页
            for page in PDFPage.create_pages(document):
                interpreter.process_page(page)
             # 接受该页面的LTPage对象
                layout=device.get_result()
                for x in layout:
                 if(isinstance(x,LTTextBoxHorizontal)):
                     with open(outfile, 'a') as f:
                         f.write(x.get_text().encode('utf-8')+'\n')
            print(sourcefile + '  已转为 ' + outfile)

##############################################把doc转为txt##############################################
# 调用之前要确保你在linux 下装了catdoc
def d2t(sourcefile, outfile):
    try:
        f =open(outfile, 'w+')
        content = os.popen("catdoc -s='utf-8' -d='utf-8' " + sourcefile).read()
        f.write(content)
        f.close()
    except:
        print(sourcefile + ' :doc文件转换失败')
    else:
        print(sourcefile + '已转为 ' + outfile)

        
############################################### 把docx转为txt#############################################
def dx2t(sourcefile, outfile):
    import docx
    try:
        fullText = []
        doc = docx.Document(sourcefile)
        f = open(outfile, 'w+')
        paras = doc.paragraphs
        for p in paras:
            fullText.append(p.text)
        f.write('\n'.join(fullText))
        f.close()
    except:
        print(sourcefile + ' :doc文件转换失败')
    else:
        print(sourcefile + '已转为 ' + outfile)